import os
import uuid  
from werkzeug.utils import secure_filename
from flask import Flask, render_template, redirect, url_for, flash, request, send_from_directory, abort, send_file
from flask_login import login_user, logout_user, login_required, current_user
from datetime import datetime

from docx import Document as DocxDoc
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

from config import Config
from models import db, User, Document, Knowledge, Collection, CollectionItem, RecentlyViewed
from forms import RegistrationForm, LoginForm, DocumentForm, DocumentEditForm
from utils import init_search_index, search_fulltext, index_document, delete_document_from_index, backup_database

def create_app(config_class=Config):
    # Запускаємо фласк і підтягуємо конфіги
    app = Flask(__name__)
    app.config.from_object(config_class)

    # Робимо папки для файлів і бекапів, якщо їх ще немає
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    os.makedirs('backups', exist_ok=True)
    os.makedirs(app.instance_path, exist_ok=True)

    # Підключаємо базу і захист форм
    db.init_app(app)
    from flask_wtf.csrf import CSRFProtect
    CSRFProtect(app)

    # Налаштовуємо вхід користувачів
    from flask_login import LoginManager
    login_manager = LoginManager(app)
    login_manager.login_view = 'login'
    login_manager.login_message = 'Будь ласка, увійдіть в систему'
    login_manager.login_message_category = 'info'

    @login_manager.user_loader
    def load_user(user_id):
        return User.query.get(int(user_id))

    # Ця штука потрібна, щоб на будь-якій сторінці показувати "Нещодавно переглянуті"
    @app.context_processor
    def inject_recent():
        if current_user.is_authenticated:
            # Беремо останні 20 переглядів, фільтруємо дублікати, залишаємо 10 унікальних
            recent_views = RecentlyViewed.query.filter_by(user_id=current_user.id)\
                .order_by(RecentlyViewed.viewed_at.desc()).limit(20).all()
            unique_recent = []
            seen = set()
            for view in recent_views:
                if view.document_id not in seen:
                    unique_recent.append(view.document)
                    seen.add(view.document_id)
                if len(unique_recent) >= 10: break
            return dict(recently_viewed_docs=unique_recent)
        return dict(recently_viewed_docs=[])

    # При старті створюємо таблиці і адміна, якщо його ще нема
    with app.app_context():
        db.create_all()
        init_search_index()
        if not User.query.filter_by(email='admin@example.com').first():
            admin = User(email='admin@example.com', name='Адміністратор системи', role='admin', is_active=True)
            admin.set_password('admin123')
            db.session.add(admin)
            db.session.commit()

    # === Сторінки входу/реєстрації ===
    @app.route('/')
    def index():
        # На головній показуємо останні завантажені документи
        recent_uploads = Document.query.order_by(Document.uploaded_at.desc()).limit(10).all()
        return render_template('index.html', recent=recent_uploads)

    @app.route('/register', methods=['GET', 'POST'])
    def register():
        if current_user.is_authenticated: return redirect(url_for('index'))
        form = RegistrationForm()
        if form.validate_on_submit():
            user = User(email=form.email.data, name=form.name.data, role='user')
            user.set_password(form.password.data)
            db.session.add(user)
            db.session.commit()
            flash('Реєстрація успішна!', 'success')
            return redirect(url_for('login'))
        return render_template('auth/register.html', form=form)

    @app.route('/login', methods=['GET', 'POST'])
    def login():
        if current_user.is_authenticated: return redirect(url_for('index'))
        form = LoginForm()
        if form.validate_on_submit():
            user = User.query.filter_by(email=form.email.data).first()
            if user and user.check_password(form.password.data) and user.is_active:
                login_user(user, remember=True)
                next_page = request.args.get('next')
                return redirect(next_page) if next_page else redirect(url_for('index'))
            flash('Невірний логін або пароль', 'danger')
        return render_template('auth/login.html', form=form)

    @app.route('/logout')
    @login_required
    def logout():
        logout_user()
        flash('Ви вийшли з системи', 'info')
        return redirect(url_for('index'))

    # === Робота з документами ===
    @app.route('/document/upload', methods=['GET', 'POST'])
    @login_required
    def upload_document():
        form = DocumentForm()
        if form.validate_on_submit():
            f = form.file.data
            original_filename = f.filename
            ext = os.path.splitext(original_filename)[1].lower()
            
            # Дозволяємо тільки PDF та DOCX
            if ext not in ['.pdf', '.docx']:
                flash('Тільки PDF та DOCX', 'danger')
                return redirect(request.url)
            
            # Зберігаємо файл з унікальним ім'ям, щоб не було конфліктів
            unique_filename = str(uuid.uuid4()) + ext
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
            f.save(filepath)
            
            doc = Document(
                title=form.title.data, authors=form.authors.data, year=form.year.data,
                source=form.source.data, doc_type=form.doc_type.data,
                original_filename=original_filename, stored_filename=unique_filename,
                uploaded_by=current_user.id
            )
            db.session.add(doc)
            db.session.commit()
            
            # Додаємо текст документа в пошуковий індекс
            index_document(doc.id, filepath)
            flash('Документ завантажено!', 'success')
            return redirect(url_for('document_list'))
        return render_template('document/upload.html', form=form)

    @app.route('/document/<int:doc_id>/edit', methods=['GET', 'POST'])
    @login_required
    def edit_document(doc_id):
        doc = Document.query.get_or_404(doc_id)
        # Редагувати може тільки адмін або автор
        if current_user.role != 'admin' and doc.uploaded_by != current_user.id: abort(403)
        
        form = DocumentEditForm(obj=doc)
        if form.validate_on_submit():
            form.populate_obj(doc)
            # Якщо завантажили новий файл — замінюємо старий
            if form.file.data:
                f = form.file.data
                ext = os.path.splitext(f.filename)[1].lower()
                new_filename = str(uuid.uuid4()) + ext
                f.save(os.path.join(app.config['UPLOAD_FOLDER'], new_filename))
                try: os.remove(os.path.join(app.config['UPLOAD_FOLDER'], doc.stored_filename))
                except: pass
                doc.stored_filename = new_filename
                doc.original_filename = f.filename
            
            db.session.commit()
            # Оновлюємо пошуковий індекс
            index_document(doc.id, os.path.join(app.config['UPLOAD_FOLDER'], doc.stored_filename))
            flash('Документ оновлено', 'success')
            return redirect(url_for('document_detail', doc_id=doc.id))
        return render_template('document/upload.html', form=form, title="Редагування")

    @app.route('/document/<int:doc_id>/delete', methods=['POST'])
    @login_required
    def delete_document(doc_id):
        doc = Document.query.get_or_404(doc_id)
        if current_user.role != 'admin' and doc.uploaded_by != current_user.id: abort(403)
        
        # Видаляємо з пошуку, з диска і з бази
        delete_document_from_index(doc.id)
        try: os.remove(os.path.join(app.config['UPLOAD_FOLDER'], doc.stored_filename))
        except: pass
        
        # Чистимо пов'язані дані (знання, історію)
        Knowledge.query.filter_by(document_id=doc.id).delete()
        RecentlyViewed.query.filter_by(document_id=doc.id).delete()
        db.session.delete(doc)
        db.session.commit()
        flash('Документ видалено', 'success')
        return redirect(url_for('document_list'))

    @app.route('/documents')
    @login_required
    def document_list():
        # Отримуємо параметри для пошуку та фільтрації
        query = request.args.get('q', '').strip()
        author = request.args.get('author', '')
        year_from = request.args.get('year_from', type=int)
        year_to = request.args.get('year_to', type=int)
        doc_type = request.args.get('type', '')
        show_my = request.args.get('show_my')

        docs = Document.query
        # Якщо є пошуковий запит — шукаємо по тексту
        if query:
            ids = search_fulltext(query)
            docs = docs.filter(Document.id.in_(ids)) if ids else docs.filter(False)
        
        # Застосовуємо фільтри
        if author: docs = docs.filter(Document.authors.ilike(f'%{author}%'))
        if year_from: docs = docs.filter(Document.year >= year_from)
        if year_to: docs = docs.filter(Document.year <= year_to)
        if doc_type: docs = docs.filter_by(doc_type=doc_type)
        if show_my == '1': docs = docs.filter_by(uploaded_by=current_user.id)
        
        documents = docs.order_by(Document.uploaded_at.desc()).all()
        return render_template('document/list.html', documents=documents)

    @app.route('/document/<int:doc_id>')
    @login_required
    def document_detail(doc_id):
        doc = Document.query.get_or_404(doc_id)
        # Фіксуємо перегляд в історії
        view = RecentlyViewed(user_id=current_user.id, document_id=doc_id)
        db.session.add(view)
        db.session.commit()
        
        knowledges = Knowledge.query.filter_by(document_id=doc_id, user_id=current_user.id).all()
        
        # Сам файл будемо показувати через JS на клієнті
        return render_template('document/detail.html', doc=doc, knowledges=knowledges)

    # Цей маршрут віддає файл, щоб його можна було переглянути в браузері
    @app.route('/document/<int:doc_id>/view')
    @login_required
    def view_document_file(doc_id):
        doc = Document.query.get_or_404(doc_id)
        return send_from_directory(app.config['UPLOAD_FOLDER'], doc.stored_filename,
                                   as_attachment=False, download_name=doc.original_filename)

    # А цей — щоб скачати файл
    @app.route('/document/<int:doc_id>/download')
    @login_required
    def download_document(doc_id):
        doc = Document.query.get_or_404(doc_id)
        return send_from_directory(app.config['UPLOAD_FOLDER'], doc.stored_filename,
                                   as_attachment=True, download_name=doc.original_filename)

    # === Знання, Нотатки та Колекції ===
    @app.route('/knowledge/add/<int:doc_id>', methods=['POST'])
    @login_required
    def add_knowledge(doc_id):
        text = request.form.get('text', '').strip()
        note = request.form.get('note', '')
        tags = request.form.get('tags', '')
        if text:
            k = Knowledge(document_id=doc_id, user_id=current_user.id, text=text, note=note, tags=tags)
            db.session.add(k)
            db.session.commit()
            flash('Збережено!', 'success')
        return redirect(url_for('document_detail', doc_id=doc_id))

    @app.route('/knowledge/<int:k_id>/edit', methods=['POST'])
    @login_required
    def edit_knowledge(k_id):
        k = Knowledge.query.get_or_404(k_id)
        if k.user_id != current_user.id: abort(403)
        
        new_text = request.form.get('text', '').strip()
        if new_text:
            k.text = new_text
            k.note = request.form.get('note', k.note)
            k.tags = request.form.get('tags', k.tags)
            db.session.commit()
            flash('Оновлено', 'success')
        else:
            flash('Текст пустий', 'danger')
        
        next_page = request.args.get('next')
        if next_page: return redirect(next_page)
        return redirect(url_for('my_knowledge'))

    @app.route('/knowledge/<int:k_id>/delete', methods=['POST'])
    @login_required
    def delete_knowledge(k_id):
        k = Knowledge.query.get_or_404(k_id)
        if k.user_id != current_user.id: abort(403)
        db.session.delete(k)
        db.session.commit()
        flash('Видалено', 'success')
        next_page = request.args.get('next')
        if next_page: return redirect(next_page)
        return redirect(url_for('my_knowledge'))

    # Головна сторінка для роботи з нотатками і колекціями
    @app.route('/my/knowledge', methods=['GET', 'POST'])
    @login_required
    def my_knowledge():
        # Збираємо купу фільтрів з GET-запиту
        q = request.args.get('q', '').strip()
        tag = request.args.get('tag', '').strip()
        filter_doc = request.args.get('filter_doc', type=int)
        filter_col = request.args.get('filter_col', type=int)
        sort_by = request.args.get('sort_by', 'date_desc')

        col_q = request.args.get('col_q', '').strip()
        col_item_q = request.args.get('col_item_q', '').strip()
        col_sort = request.args.get('col_sort', 'name_asc')
        active_tab = request.args.get('tab', 'all')

        # Будуємо запит на нотатки
        query = Knowledge.query.filter_by(user_id=current_user.id)
        if q: query = query.filter((Knowledge.text.ilike(f'%{q}%')) | (Knowledge.note.ilike(f'%{q}%')))
        if tag: query = query.filter(Knowledge.tags.ilike(f'%{tag}%'))
        if filter_doc: query = query.filter_by(document_id=filter_doc)
        if filter_col: query = query.join(Knowledge.collection_items).filter(CollectionItem.collection_id == filter_col)

        # Сортуємо нотатки
        if sort_by == 'doc_title': query = query.join(Document).order_by(Document.title.asc())
        elif sort_by == 'date_asc': query = query.order_by(Knowledge.created_at.asc())
        else: query = query.order_by(Knowledge.created_at.desc())
        knowledges = query.all()

        # Тепер розбираємося з колекціями
        c_query = Collection.query.filter_by(user_id=current_user.id)
        if col_q: c_query = c_query.filter(Collection.name.ilike(f'%{col_q}%'))
        if col_item_q: c_query = c_query.join(Collection.items).join(CollectionItem.knowledge)\
                             .filter((Knowledge.text.ilike(f'%{col_item_q}%')) | (Knowledge.note.ilike(f'%{col_item_q}%')))
        
        collections_list = c_query.all()
        # Сортування колекцій робимо в пайтоні
        if col_sort == 'count_desc': collections_list.sort(key=lambda c: len(c.items), reverse=True)
        elif col_sort == 'count_asc': collections_list.sort(key=lambda c: len(c.items), reverse=False)
        elif col_sort == 'name_desc': collections_list.sort(key=lambda c: c.name.lower(), reverse=True)
        else: collections_list.sort(key=lambda c: c.name.lower())

        # Для випадаючих списків у фільтрах
        all_docs = Document.query.join(Knowledge).filter(Knowledge.user_id==current_user.id).distinct().all()
        all_cols = Collection.query.filter_by(user_id=current_user.id).order_by(Collection.name).all()

        # Обробка масових дій (експорт або додавання в колекцію)
        if request.method == 'POST':
            action = request.form.get('action')
            
            # Експорт вибраних нотаток у DOCX
            if action == 'export_docx':
                selected_ids = request.form.getlist('knowledge_ids')
                ordered_ids_str = request.form.get('ordered_ids', '')
                if not selected_ids:
                    flash('Нічого не вибрано', 'warning')
                    return redirect(url_for('my_knowledge'))
                
                # Визначаємо порядок записів (якщо юзер їх перетягував)
                final_ids = []
                if ordered_ids_str:
                    click_order = ordered_ids_str.split(',')
                    for oid in click_order:
                        if oid in selected_ids: final_ids.append(int(oid))
                    for sid in selected_ids:
                        if int(sid) not in final_ids: final_ids.append(int(sid))
                else: final_ids = [int(x) for x in selected_ids]

                k_objects = Knowledge.query.filter(Knowledge.id.in_(final_ids)).all()
                k_map = {k.id: k for k in k_objects}
                sorted_knowledge = [k_map[fid] for fid in final_ids if fid in k_map]

                # Формуємо документ
                doc = DocxDoc()
                doc.add_heading('Експортовані конспекти', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
                for i, k in enumerate(sorted_knowledge, 1):
                    doc.add_heading(f"{i}. {k.document.title}", level=1)
                    doc.add_paragraph(f"Автори: {k.document.authors}")
                    doc.add_paragraph(k.text, style='Intense Quote')
                    if k.note:
                        p = doc.add_paragraph()
                        p.add_run("Примітка: ").bold = True
                        p.add_run(k.note)
                    doc.add_paragraph()
                
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return send_file(buffer, as_attachment=True, download_name=f"export_{datetime.now().strftime('%H-%M')}.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

            # Додавання вибраного до колекції
            elif action == 'add_to_collection':
                collection_id = request.form.get('collection_id')
                selected_ids = request.form.getlist('knowledge_ids')
                if collection_id and selected_ids:
                    coll = Collection.query.get(collection_id)
                    if coll and coll.user_id == current_user.id:
                        count = 0
                        for kid in selected_ids:
                            kn = Knowledge.query.get(int(kid))
                            exists = CollectionItem.query.filter_by(collection_id=coll.id, knowledge_id=kn.id).first()
                            if kn and kn.user_id == current_user.id and not exists:
                                item = CollectionItem(collection=coll, knowledge=kn)
                                db.session.add(item)
                                count += 1
                        db.session.commit()
                        flash(f'Додано {count} записів', 'success')
                return redirect(url_for('my_knowledge'))

        return render_template('knowledge/list.html', knowledges=knowledges, collections=collections_list, all_docs=all_docs, all_cols=all_cols, active_tab=active_tab)

    # === Управління колекціями ===
    @app.route('/collection/create', methods=['POST'])
    @login_required
    def create_collection():
        name = request.form.get('name', '').strip()
        if name:
            if not Collection.query.filter_by(name=name, user_id=current_user.id).first():
                db.session.add(Collection(name=name, user_id=current_user.id))
                db.session.commit()
            else: flash('Вже існує', 'warning')
        return redirect(url_for('my_knowledge', tab='collections'))

    @app.route('/collection/<int:c_id>/delete', methods=['POST'])
    @login_required
    def delete_collection(c_id):
        c = Collection.query.get_or_404(c_id)
        if c.user_id != current_user.id: abort(403)
        db.session.delete(c)
        db.session.commit()
        return redirect(url_for('my_knowledge', tab='collections'))
    
    @app.route('/collection/<int:c_id>/rename', methods=['POST'])
    @login_required
    def rename_collection(c_id):
        c = Collection.query.get_or_404(c_id)
        if c.user_id != current_user.id: abort(403)
        c.name = request.form.get('name', c.name)
        db.session.commit()
        return redirect(url_for('my_knowledge', tab='collections'))

    @app.route('/collection/<int:c_id>/remove_item/<int:k_id>', methods=['POST'])
    @login_required
    def remove_from_collection(c_id, k_id):
        c = Collection.query.get_or_404(c_id)
        if c.user_id != current_user.id: abort(403)
        item = CollectionItem.query.filter_by(collection_id=c.id, knowledge_id=k_id).first()
        if item:
            db.session.delete(item)
            db.session.commit()
        return redirect(url_for('my_knowledge', tab='collections'))

    @app.route('/collection/<int:c_id>/export/docx')
    @login_required
    def export_collection_docx(c_id):
        # Експорт цілої колекції в DOCX
        c = Collection.query.get_or_404(c_id)
        if c.user_id != current_user.id: abort(403)
        doc = DocxDoc()
        doc.add_heading(f'Колекція: {c.name}', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, item in enumerate(c.items, 1):
            k = item.knowledge
            doc.add_heading(f"{i}. {k.document.title}", level=1)
            doc.add_paragraph(f"Автори: {k.document.authors}")
            doc.add_paragraph(k.text, style='Intense Quote')
            if k.note:
                p = doc.add_paragraph()
                p.add_run("Примітка: ").bold = True
                p.add_run(k.note)
            doc.add_paragraph()
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name=f"{secure_filename(c.name)}.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    # === Адмінка ===
    @app.route('/admin/users')
    @login_required
    def admin_users():
        if current_user.role != 'admin': abort(403)
        users = User.query.all()
        return render_template('admin/users.html', users=users)

    @app.route('/admin/user/<int:user_id>/toggle', methods=['POST'])
    @login_required
    def toggle_user(user_id):
        # Блокування/розблокування юзера
        if current_user.role != 'admin': abort(403)
        if user_id == current_user.id: return redirect(url_for('admin_users'))
        user = User.query.get_or_404(user_id)
        user.is_active = not user.is_active
        db.session.commit()
        return redirect(url_for('admin_users'))

    @app.route('/admin/backup')
    @login_required
    def admin_backup():
        if current_user.role != 'admin': abort(403)
        path = backup_database()
        flash(f'Бекап: {os.path.basename(path)}', 'success')
        return redirect(url_for('index'))

    return app

if __name__ == '__main__':
    app = create_app()
    app.run(debug=True)